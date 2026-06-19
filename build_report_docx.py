from pathlib import Path
from docx import Document


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paras(doc, text):
    for block in text.strip().split("\n\n"):
        doc.add_paragraph(block.strip())


def main():
    root = Path(__file__).parent
    out = root / "KANEM_Waste_Project_Report_Draft.docx"
    ch3_path = root / "CHAPTER_THREE_DESIGN.md"

    doc = Document()
    add_heading(doc, "INTELLIGENT WASTE MANAGEMENT SYSTEM", 0)
    doc.add_paragraph(
        "Compiled draft containing Chapters 1-5, references, and appendices."
    )

    add_heading(doc, "CHAPTER ONE", 1)
    add_heading(doc, "1.3 Research Questions", 2)
    add_paras(
        doc,
        """
How can a web-based platform be designed to improve waste pickup request submission and interaction?
In what ways can cloud-based storage improve the accuracy and accessibility of records?
How can AI-assisted guidance support better waste description and classification before pickup?
How can an administrative dashboard improve monitoring and status management of requests?
What communication workflow can provide timely updates to users?
To what extent does the system improve efficiency, usability, and service responsiveness compared with manual methods?
""",
    )
    add_heading(doc, "1.8 Definition of Operational Terms", 2)
    add_paras(
        doc,
        """
Operational terms include smart waste management system, AI, computer vision, web application, cloud database,
waste pickup request, route optimization, TSP, administrative dashboard, notification workflow, usability, and
performance efficiency. These are interpreted according to their direct practical usage in this project context.
""",
    )

    add_heading(doc, "CHAPTER TWO", 1)
    add_heading(doc, "2.1 Conceptual Framework", 2)
    add_paras(
        doc,
        """
The conceptual framework links user-facing web workflows, AI-assisted guidance, cloud data management,
admin monitoring, and route optimization to improved service quality and operational efficiency.
""",
    )
    add_heading(doc, "2.2 Related Works and Research Gaps", 2)
    add_paras(
        doc,
        """
Related works confirm the value of AI, IoT, and optimization in waste management, but many are model-centric
or smart-city oriented. The major gap addressed here is a lightweight, deployable, business-level integrated platform.
""",
    )
    add_heading(doc, "2.3 Theoretical Framework", 2)
    add_paras(
        doc,
        """
This study is grounded in Systems Theory, TAM, TPB, TSP optimization principles, and usability-centered design.
These theories jointly explain technical integration, user adoption, and route efficiency outcomes.
""",
    )
    add_heading(doc, "2.4 Summary of the Review", 2)
    add_paras(
        doc,
        """
The review supports feasibility, reveals deployment gaps in existing studies, and justifies this integrated
AI-enabled web approach for Kanem Waste Solutions.
""",
    )

    add_heading(doc, "CHAPTER THREE", 1)
    add_heading(doc, "3.0 System Analysis and Design", 2)
    add_paras(
        doc,
        """
System analysis identified limitations of the existing manual process and justified an integrated digital solution.
Detailed design artifacts (ERD, use case, deployment, class, and activity diagrams) are included in Chapter Three.
""",
    )
    if ch3_path.exists():
        add_heading(doc, "Chapter Three Detailed Design Artifacts", 2)
        text = ch3_path.read_text(encoding="utf-8", errors="ignore")
        # keep as plain paragraphs for Word readability
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("#"):
                add_heading(doc, line.lstrip("# ").strip(), 3)
            else:
                doc.add_paragraph(line)

    add_heading(doc, "CHAPTER FOUR", 1)
    add_heading(doc, "4.1 System/Network Requirement for Development", 2)
    add_paras(
        doc,
        """
The system requires a modern development machine, Node.js runtime, React frontend stack, cloud database access,
secure backend API handling, and stable internet connectivity for cloud and AI service integrations.
""",
    )
    add_heading(doc, "4.2 System Menus Implementation (Frontend)", 2)
    add_paras(
        doc,
        """
Core interfaces implemented include public navigation, pickup request flow, AI guidance entry, admin dashboard tabs
(pickups, volunteers, route optimizer), responsive mobile drawer navigation, status modals, filtering, and pagination.
""",
    )
    add_heading(doc, "4.3 Database Implementation (Backend)", 2)
    add_paras(
        doc,
        """
Backend database implementation uses cloud storage for pickup and volunteer records with real-time synchronization,
status lifecycle persistence, and operational retrieval for dashboard workflows.
""",
    )
    add_heading(doc, "4.4 System Testing", 2)
    add_paras(
        doc,
        """
Functional tests validated request creation, filtering, status updates, optimization interactions, and modal actions.
Non-functional evaluation used usability criteria (learnability, efficiency, satisfaction), responsiveness, and basic performance checks.
""",
    )
    add_heading(doc, "4.5 Performance Evaluation (Optional)", 2)
    add_paras(
        doc,
        """
Evaluation metrics include AI guidance utility, route distance/time reduction, response and completion rates,
and notification delivery effectiveness.
""",
    )

    add_heading(doc, "CHAPTER FIVE", 1)
    add_heading(doc, "5.1 Summary", 2)
    add_paras(
        doc,
        """
The project successfully delivered an AI-enabled, cloud-backed, responsive waste management platform that improves
request handling, operational monitoring, and service communication for Kanem Waste Solutions.
""",
    )
    add_heading(doc, "5.2 Conclusion", 2)
    add_paras(
        doc,
        """
The system is technically feasible and operationally beneficial, providing a practical alternative to manual methods
while supporting scalable digital transformation in local waste services.
""",
    )
    add_heading(doc, "5.3 Recommendations", 2)
    add_paras(
        doc,
        """
Future improvements should include stronger security hardening, richer analytics, advanced route constraints,
expanded notification channels, broader pilot deployment, and continuous AI quality refinement.
""",
    )

    add_heading(doc, "REFERENCES", 1)
    refs = [
        "Alourani, A., Ashraf, M. U., & Aloraini, M. (2025). Smart waste management and classification system using advanced IoT and AI technologies. PeerJ Computer Science, 11, e2777.",
        "Anitha, R. A. R., & Parthiban, A. P. A. (2025). AI-IoT-graph synergy for smart waste management. Frontiers in Sustainability, 6, 1675021.",
        "Shahab, S., Anjum, M., & Umar, M. S. (2022). Deep learning applications in solid waste management. IJACSA, 13(3).",
        "Kaza, S., Yao, L., Bhada-Tata, P., & Van Woerden, F. (2018). What a Waste 2.0. World Bank.",
        "Nielsen, J. (1994). Usability Engineering. Morgan Kaufmann.",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    add_heading(doc, "APPENDIX (A)", 1)
    add_paras(doc, "Interface specifications, sample forms, and test evidence templates.")
    add_heading(doc, "APPENDIX (B)", 1)
    add_paras(doc, "Source code module index and implementation snapshot references.")
    add_heading(doc, "APPENDIX (C)", 1)
    add_paras(doc, "User manual, installation guide, and operations checklist.")

    doc.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    main()

